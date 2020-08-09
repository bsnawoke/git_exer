# -*- coding: utf-8 -*-

'''
word table for STD document
'''

import docx
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Length
from docx.oxml.ns import qn

from docx.enum.text import WD_LINE_SPACING # 行间距
from docx.enum.text import WD_ALIGN_PARAGRAPH # 段落样式
from docx.enum.table import WD_TABLE_ALIGNMENT # 表格样式
from docx.enum.table import WD_ALIGN_VERTICAL

import docx.package
import docx.parts.document
import docx.parts.numbering




def set_col_width(table):
    widths = [ Cm(3), Cm(5), Cm(3), Cm(5) ]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width



location = 'C:/Users/bsnawoke/Desktop/table_test.docx'

file = docx.Document(location)
#file = docx.Document()

file.styles['Normal'].font.name = 'Times New Roman'
file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
name_list =[
    "case_name",
    "case_id",
    "case_descri",
    "case_cdtion",
    "case_termin",
    "case_expect",
    "case_step",
    "case_result"
    ]
case ={
       "case_name" : "wdnmd",
       "case_id" : "CT_ST_ZT_AT",
       "case_descri" : "描述",
       "case_cdtion" : "前提",
       "case_termin" : "终止",
       "case_expect" : "\n将测试aaa",
       "case_step" : "None",
       "case_result" : "预期结果"
       }

###
# 样式列表:标题:x级标题 x=1,2,3,4,5
#         图片:图片-图片标题
#         表格:表格标题-表格首行-表格内容 Table_Standard
#         注释、正文  
### 

heading = file.add_heading(text=case['case_name'],level=2)
table = file.add_table(rows=6,cols=4,style="Table_Standard")
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_col_width(table)


table.cell(0,0).text = "用例名称"
table.cell(0,1).text = case['case_name']
table.cell(0,2).text = "用例标识"
table.cell(0,3).text = case['case_id']
table.cell(0,3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
table.cell(0,3).vertical_alignment= WD_ALIGN_VERTICAL.CENTER
# 表格+图片
table.cell(4,1).paragraphs[0].add_run().add_picture('E:/entertainments/pictures/07.jpg')
#table.cell(0,3).line_spacint_rule = WD_LINE_SPACING.MULTIPLE


for i in range(1,6):
    table.cell(i,1).merge(table.cell(i,3))
    table.cell(i,1).text = case[name_list[i+1]]
    #table.cell(i,1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    table.cell(i,1).width = Inches(2)
    
para = file.add_paragraph('faq')
# 首行缩进
para.paragraph_format.first_line_indent = Cm(0.84)
#para.alignment = WD_ALIGN_PARAGRAPH.CENTER
file.add_table(1,1)
para.add_run("dddd")

#file.add_picture('E:/entertainments/pictures/07.jpg')
#file.add_paragraph("wdnmd",style="List Number")

para_a = file.add_paragraph("阿这",style="一级标题")
para_b = file.add_paragraph("就这?",style="二级标题")
para_c = file.add_paragraph("哦这样啊",style="三级标题")
para_a.add_run("wtf")
para_b.add_run("nmsl")

for para in file.paragraphs:
    print(para.text)
    if para.text == "插入位置":
        pass


try:    
    file.save(location)
except PermissionError as e :
    print(f"{e}:文件已打开")
finally:
    print('done')
